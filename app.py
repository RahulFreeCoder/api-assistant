import smtplib
import ssl
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import uvicorn
from sentence_transformers import SentenceTransformer, util
import json
import os
import requests
import glob
import numpy as np
import re
import torch
from io import BytesIO
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, File, UploadFile

load_dotenv(override=True)
# Does not need pushover text as of now
# def push(text):
#     requests.post(
#         "https://api.pushover.net/1/messages.json",
#         data={
#             "token": os.getenv("PUSHOVER_TOKEN"),
#             "user": os.getenv("PUSHOVER_USER"),
#             "message": text,
#         }
#     )
# print text to console

# 1️⃣ Predefined skill dictionary
SKILLS_DICT = [
    "Frontend/UI" "Micro-Frontend", "React.js", "Angular", "HTML5", "CSS3", "Bootstrap", "Karma/Jasmine/Jest", "AG grid",
    "Java", "J2EE", "Spring Boot", "REST/SOAP", "Reactive Microservices", "JPA", "Hibernate",
    "Rsocket", "Scala", "nodejs", "Redis", "Kafka", "Backend & APIs", "SRE & Observability", "Prometheus", "Grafana",
    "Micrometer", "Zipkin", "Sleuth", "Alerting Frameworks",
    "Reliability Testing", "SLI/SLO/SLA governance",
    "Cloud & DevOps", "Lambda", "API Gateway", "Cognito", "SNS", "SES", "SQS", "IAM", "GCP", "OpenShift",
    "Kubernetes", "Docker", "Jenkins", "Maven", "Gradle", "Github Actions", "Jenkins", "Github actions",
    "Architecture", "Microservices patterns", "DDD", "Event-Driven Architecture", "Clean Architecture", "TDD",
    "Design Patterns",
    "Database", "Oracle", "PostgreSQL", "DynamoDB", "DB2", "MySQL", "Mongodb",
    "Testing", "JUnit", "TestNG", "Cucumber", "Loadrunner", "Apache jmeter", "Gattling",
    "Python", "React", "ReactJS", "ML", "NLP", "Cloud",
    "AWS", "GCP", "Docker", "Kubernetes", "TensorFlow", "PyTorch",
]

def push(text):
    print(text)
# send email uing Gmail SMTP
def send_email(to_email, subject, body):
    smtp_server = "smtp.gmail.com"
    smtp_port = 587
    sender_email = os.getenv("EMAIL_USER")   # your Gmail
    sender_pass = os.getenv("EMAIL_PASS")   # app password

    msg = MIMEText(body)
    msg["Subject"] = subject
    msg["From"] = sender_email
    msg["To"] = to_email
    print("send email")
    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(sender_email, sender_pass)
        server.sendmail(sender_email, [to_email], msg.as_string())

def extract_skills_jd(jd_text):
    """
    Extract skills from JD text.

    Args:
        jd_text (str): Job description text
        use_llm_fallback (bool): If True, calls Gemini LLM for skill extraction

    Returns:
        list: list of extracted skills
    """
    use_llm_fallback = 1
    # --- Rule-based extraction ---
    rule_skills = [skill for skill in SKILLS_DICT if re.search(rf"\b{skill}\b", jd_text, re.IGNORECASE)]
    # --- Optional: LLM fallback ---
    if use_llm_fallback:
       skills_json = me.extract_skills_list(jd_text)
    
     # Check the type of skills_json before processing
    if isinstance(skills_json, str):
        # If it's a string, split it
        skills = [s.strip() for s in skills_json.split(",") if s.strip()]
    elif isinstance(skills_json, list):
        # If it's already a list, use it directly
        skills = skills_json
    else:
        # Handle any other unexpected types
        skills = [] # or handle as an error

    return skills
    
        

def record_user_details(email, name="Name not provided", notes="not provided"):
    body = f"Recruiter shared details:\n\nName: {name}\nEmail: {email}\nNotes: {notes}"
    subject = "Recruiter Interested!"
    body = f"A recruiter shared their email: {email}\n\nReply ASAP!"
    send_email("rahul.d.kolhe@gmail.com", subject, body)
    return {"recorded": "ok"}

def record_unknown_question(question):
    push(f"Recording {question}")
    return {"recorded": "ok"}

# A simple function to generate the suggestion text.
def suggest_job_match_analyzer():
    """
    Suggests using the Job Match Analyzer to the user.
    """
    return {
        "suggestion": "I can help with that! I have a Job Match Analyzer tool that compares your skills to a job description. Please upload your resume and the job description, and I'll give you a detailed report on how well you match."
    }

record_user_details_json = {
    "name": "record_user_details",
    "description": "Use this tool to record that a user is interested in being in touch and provided an email address",
    "parameters": {
        "type": "object",
        "properties": {
            "email": {
                "type": "string",
                "description": "The email address of this user"
            },
            "name": {
                "type": "string",
                "description": "The user's name, if they provided it"
            }
            ,
            "notes": {
                "type": "string",
                "description": "Any additional information about the conversation that's worth recording to give context"
            }
        },
        "required": ["email"],
        "additionalProperties": False
    }
}

record_unknown_question_json = {
    "name": "record_unknown_question",
    "description": "Always use this tool to record any question that couldn't be answered as you didn't know the answer or user is asking more than 3 question",
    "parameters": {
        "type": "object",
        "properties": {
            "question": {
                "type": "string",
                "description": "The question that couldn't be answered"
            },
        },
        "required": ["question"],
        "additionalProperties": False
    }
}
# The JSON schema for the new tool.
suggest_job_match_analyzer_json = {
    "name": "suggest_job_match_analyzer",
    "description": "Use this tool to proactively suggest the Job Match Analyzer tab on current site to the user when they ask about matching their resume to a job or finding a good fit.",
    "parameters": {
        "type": "object",
        "properties": {},
        "additionalProperties": False
    }
}

tools = [{"type": "function", "function": record_user_details_json},
        {"type": "function", "function": record_unknown_question_json},
        {"type": "function", "function": suggest_job_match_analyzer_json}]

#FAST API 

origins = [
        "http://localhost:3000",   # CRA dev server
        "http://127.0.0.1:3000",
        "http://localhost:5173",   # Vite dev server
        "http://127.0.0.1:5173",
        "https://stately-jelly-2c936a.netlify.app/"
]
app = FastAPI()
    # Allow React frontend to call API (CORS)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    )

# --- FastAPI Endpoints ---
@app.get("/api/hello")
def hello():
    return {"message": "Hello from Python API 🚀"}

@app.post("/api/ask")
async def ask(payload: dict):
    question = payload.get("question", "")
    history = payload.get("history", [])
    answer = me.chat(question, history)
    return {"answer": answer}

@app.post("/api/match")
async def match_job(file: UploadFile = File(...)):
    # Extract JD text
    text = ""
    file_extension = file.filename.split('.')[-1].lower()

    if file_extension == "pdf":
        reader = PdfReader(file.file)
        for page in reader.pages:
            text += page.extract_text() or ""
    elif file_extension == "docx":
        docx_bytes = await file.read()
        docx_stream = BytesIO(docx_bytes)
        
        document = Document(docx_stream)
        
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
    else:
        # Handle other file types or raise an error
        # In this case, we'll just read the content as a string
        text = (await file.read()).decode("utf-8")

    print(f"print JD: {text}" )
    # Get embedding for JD
    if text is not None:
        jd_skills_list = extract_skills_jd(text)
        print(f"Extracted skills from JD: {jd_skills_list}")
        if jd_skills_list:
            jd_embedding = me.get_embedding(jd_skills_list)
        else:
            jd_embedding = None  # handle in downstream code
        
        sim_matrix = util.cos_sim(jd_embedding, me.resume_skills_embedding)  # shape: (num_jd_skills, num_profile_texts)
        max_sim_per_jd = sim_matrix.max(dim=1).values  # take highest similarity for each JD skill
        matched_skills, partial_skills, missing_skills = [], [], []
        for idx, score in enumerate(max_sim_per_jd):
            skill = jd_skills_list[idx]
            if score >= 0.75:
                matched_skills.append(skill)
            elif score >= 0.5:
                partial_skills.append(skill)
            else:
                missing_skills.append(skill)
                
    total_skills = len(jd_skills_list)
    match_percentage = round((len(matched_skills) + 0.5 * len(partial_skills)) / total_skills * 100)
    print(f" Relevence score {match_percentage}")
    #provide reasoning for missing skills
    reasoning = {}
    if missing_skills is not None:
        ctr = 0
        for skill in missing_skills:
            ctr += 1
            if ctr < 3:
                reasoning[skill] = me.get_reasoning(skill)

    output = {
    "relevance": match_percentage,
    "message": f"Your profile matches {match_percentage}% with this JD 🚀",
    "matched_skills": matched_skills,
    "partial_skills": partial_skills,
    "missing_skills": missing_skills,
    "reasoning": reasoning
}
    # Return the results
    return output

@app.post("/api/generate")
async def generate_resume(file: UploadFile = File(...)):
    # Extract JD text
    text = ""
    file_extension = file.filename.split('.')[-1].lower()

    if file_extension == "pdf":
        reader = PdfReader(file.file)
        for page in reader.pages:
            text += page.extract_text() or ""
    elif file_extension == "docx":
        docx_bytes = await file.read()
        docx_stream = BytesIO(docx_bytes)
        
        document = Document(docx_stream)
        
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
    else:
        # Handle other file types or raise an error
        # In this case, we'll just read the content as a string
        text = (await file.read()).decode("utf-8")

    print(f"print JD: {text}" )
    # Get embedding for JD
    if text is not None:
        #Analyze JD
        #Synthesis
        answer = me.generator(text)
    
    return {"answer": answer}
  
class Me:
    def __init__(self,summary_text=None):

        google_api_key = os.getenv('GOOGLE_API_KEY')   
        if google_api_key:
            print(f"Google API Key exists and begins {google_api_key[:8]}")
        else:
            print("Google API Key not set - please head to the troubleshooting guide in the setup folder")
    
        GEMINI_BASE_URL = "https://generativelanguage.googleapis.com/v1beta/openai"
        self.gemini = OpenAI(base_url=GEMINI_BASE_URL, api_key=google_api_key)
        self.model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2') 
        self.name = "Rahul Kolhe"
        # Path to your PDF files (adjust as needed)
        pdf_files = glob.glob("me/*.pdf")   # reads all PDFs inside 'pdfs' folder
        for file in pdf_files:
            reader = PdfReader(file)
      
        self.linkedin = ""
        for page in reader.pages:
            text = page.extract_text()
            if text:
                self.linkedin += text
        
        with open("me/summary.txt", "r", encoding="utf-8") as f:
            self.summary = f.read()

        # Load skills from a separate file for fine-grained comparison
        with open("me/skills.txt", "r", encoding="utf-8") as f:
            self.skills = f.read()
            self.summary = self.summary + self.skills

        resume_skills_list = self.extract_skills_list(self.summary)
        print(f"Extracted skills from resume: {resume_skills_list}")
        if resume_skills_list:
            self.resume_skills_embedding  = self.get_embedding(resume_skills_list)
        else:
            self.resume_skills_embedding  = None  # handle in downstream code
    
    def extract_skills_list(self, text):
        prompt = f"""
            Extract all the key technical and business skills mentioned in the following description. 
            Return the skills as a comma-separated list includes Technical skills (like programming languages, frameworks, tools, cloud platforms, etc.) 
            and business skills (like project management, communication, leadership, stakeholder management, etc.) 
            if possible.

            Description:
            {text}
            """
        response = self.gemini.chat.completions.create(
                model="gemini-2.5-flash-preview-05-20",
                messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
                ],
            )
            # Parse the JSON response to extract skills
        try:
            skills_json = response.choices[0].message.content
                
        except Exception as e:
                print(f"Error parsing LLM response: {e} use dictionary skills")
                # If empty, fallback to empty list or default skills
        if not skills_json:
            skills = []  # or fallback SKILLS_DICT
        else:
            skills = [s.strip() for s in skills_json.split(",") if s.strip()]        
        return skills  
    
    def get_embedding(self, text):
        """Generates an embedding for the given text."""
        if not text:
            return None # Handle empty text gracefully
        print(f"Encoding text for embedding: {text[:50]}...")
        return self.model.encode(text, convert_to_tensor=True)
    
    def get_reasoning(self, jd_text):
        """
        Analyzes JD for key skills and provide reasoning.
        """
        prompt=f"Explain missing skill reasoning concisely: {jd_text} based on this profile: {me.summary}"
        try:
            response = self.gemini.chat.completions.create(
                model="gemini-2.5-flash-preview-05-20",  # Specify a Gemini model
                messages=[
                    {"role": "system", "content": "You are a helpful assistant"},
                    {"role": "user", "content": prompt}
                ],
            )
             # Check if response exists and extract content properly
            if response.choices:
                message = response.choices[0].message
                
                if hasattr(message, 'content'):
                    response_text = getattr(message, 'content')
                    
                    if isinstance(response_text, str) and response_text.strip():
                        return [line for line in response_text.split('\n') if line]
            
            # Return empty list if no valid content
            return 'Skill Not Found'

        except Exception as e:
            # Handle potential API errors gracefully
            print(f"An error occurred during Gemini API call: {e}")
        return "An error occurred while generating reasoning."

    def get_missing_skills(self, jd_text):
        """
        Analyzes JD for key skills and compares them to resume.
        This is a simplified approach. A more robust solution might involve
        a more sophisticated NLP model or a large language model.
        """
        prompt = f"Extract a concise list of the most critical technical skills, qualifications, and core responsibilities from the following job description:\n\n{jd_text}\n\nList them one per line."
        # Use your configured OpenAI model to extract key skills
        response = self.gemini.chat.completions.create(
            model="gemini-2.5-flash-preview-05-20",  # Specify a Gemini model
            messages=[
                {"role": "system", "content": "You are a helpful assistant"},
                {"role": "user", "content": prompt}
            ]
        )
        extracted_skills = response.choices[0].message.content.strip().split('\n')
        return [skill for skill in extracted_skills if skill]

    def handle_tool_call(self, tool_calls):
        results = []
        for tool_call in tool_calls:
            tool_name = tool_call.function.name
            arguments = json.loads(tool_call.function.arguments)
            print(f"Tool called: {tool_name}", flush=True)
            tool = globals().get(tool_name)
            result = tool(**arguments) if tool else {}
            results.append({"role": "tool","content": json.dumps(result),"tool_call_id": tool_call.id})
        return results
    
    def system_prompt(self):
        system_prompt = f"You are acting as {self.name}. You are answering questions on {self.name}'s website, \
particularly questions related to {self.name}'s career, professional background, leaderhsip domain and most important technical skills and experience. \
Your responsibility is to represent {self.name} for interactions on the website as faithfully as possible. \
You are given a summary of {self.name}'s professional background and LinkedIn profile, resume which you can use to answer questions. Keep ypur answer concise and try to provide details only when asked. \
Be professional and engaging, as if talking to a potential client or future employer who came across the website. \
If you don't know the answer to any question, use your record_unknown_question tool to record the question that you couldn't answer, even if it's about something trivial or unrelated to career. \
If the user is engaging in discussion and asking more than 2 questions, try to steer them towards getting in touch via email; provide linkedin profile url; email id to connect; ask for their email and record it using your record_user_details tool.\
Your tools:\
- `record_unknown_question`: Use this tool immediately if you cannot answer a user's question, regardless of its topic.\
- `record_user_details`: If the user shows interest in connecting (e.g., asking for contact info) or has been in a long conversation (more than 2 questions), proactively offer to connect and use this tool to record their email and name.\
- `suggest_job_match_analyzer`: Proactively use this tool when a user asks for help with their resume, a job description, or asks if their skills are a good match or fit for a role. This tool's purpose is to explain the value of the Job Match Analyzer and guide the user on how to use it.\
Instructions for using tools:\
- **Prioritize suggesting the job match analyzer** if the user's query is about a resume, a job, or a skill match. This is a key feature of the website.\
- If a user asks a question you can't answer, immediately use `record_unknown_question`.\
- After answering a few questions or if the user asks for contact info, use the `record_user_details` tool.\
"

        system_prompt += f"\n\n## Summary:\n{self.summary}\n\n## LinkedIn Profile and Resume:\n{self.linkedin}\n\n"
        system_prompt += f"With this context, please chat with the user, always staying in character as {self.name}."
        return system_prompt
    
    def chat(self, message, history):
       # Convert React history into Gemini-friendly format
        formatted_history = []
        for h in history:
            role = "user" if h["sender"] == "user" else "assistant"
            formatted_history.append({"role": role, "content": h["text"]})

        messages = (
            [{"role": "system", "content": self.system_prompt()}]
            + formatted_history
            + [{"role": "user", "content": message}]
        )
        done = False
        while not done:
            response = self.gemini.chat.completions.create(model="gemini-2.5-flash-preview-05-20", messages=messages, tools=tools)
            if response.choices[0].finish_reason=="tool_calls":
                message = response.choices[0].message
                tool_calls = message.tool_calls
                results = self.handle_tool_call(tool_calls)
                messages.append(message)
                messages.extend(results)
            else:
                done = True
        return response.choices[0].message.content
    
    def generator(self, jd):
        generatorprompt = f"You are an expert resume writer and a senior technical recruiter. Your goal is to create a professional, highly-effective resume tailored for a specific job description using the candidate's provided 19 years of industry experience. Your output should be in a clear, professional, ATS-friendly format.\
**Instructions:**\
1.  **Analyze the Job Description (JD):** Read the JD and extract all essential keywords, required technical skills, soft skills (e.g., leadership, communication), and core responsibilities. Also, identify any nice-to-have or preferred skills.\
2.  **Evaluate Experience:** Read the candidate's professional history. Match their experience, skills, and accomplishments to the requirements from the JD. For each requirement, find a relevant example from the candidate's history. For senior-level roles, focus on leadership, architecture, and mentoring.\
3.  **Rewrite and Generate:** Using the analysis, generate a complete resume.\
    * **Professional Summary:** Write a concise summary (3-4 lines) that immediately grabs attention by highlighting your 19 years of experience and mentioning the most important skills from the JD.\
    * **Experience Section:** For each relevant past role, rewrite 3-5 bullet points. Start each bullet point with a strong action verb. Focus on quantifiable achievements and outcomes. Ensure you incorporate the keywords and responsibilities from the JD. For example, if the JD mentions RESTful APIs, your bullet point should explicitly mention how you worked with or led projects involving them.\
    * **Skills Section:** Create a clean, organized skills section. Categorize skills into relevant groups (e.g., Languages, Frameworks, Cloud, Databases and list all skills mentioned in the JD that are present in your experience.\
**Candidate's Raw Experience:**\
{self.summary} + {self.linkedin}\
**Target Job Description:**\
{jd}\
**Final Output:**\
[Generate the complete, tailored resume in a single, well-formatted block of text.]"
        
        messages = (
            [{"role": "system", "content": generatorprompt}]

            + [{"role": "user", "content": jd}]
        )
        response = self.gemini.chat.completions.create(model="gemini-2.5-flash-preview-05-20", messages=messages)

        return response.choices[0].message.content            
    

me = Me()

if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=False)

    